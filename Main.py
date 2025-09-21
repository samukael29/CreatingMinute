from pathlib import Path
import subprocess
import os
from moviepy import VideoFileClip
import openai
import whisper
from docx import Document


def mkv_to_mp3(directory: str, filename: str):
    """
    Extrai o áudio de um arquivo MKV e salva como MP3 usando ffmpeg.
    É necessário ter o ffmpeg instalado no sistema.
    """

    inputfile = os.path.join(directory,  filename + "." + 'mkv')
    outputfile = os.path.join(directory,  filename + "." + 'mp3')

    if os.path.isfile(inputfile):
        print(f"The file '{inputfile}' exists.")
    else:
        print(f"The file '{inputfile}' does not exist or is not a regular file.")

    command = [
        "ffmpeg",
        "-i", inputfile,   # arquivo de entrada
        "-q:a", "0",        # qualidade máxima de áudio
        "-map", "a",        # seleciona apenas a trilha de áudio
        outputfile
    ]

    subprocess.run(command, check=True)
    print(f"Áudio exportado com sucesso para: {outputfile}")



def transcribe_mp3_to_text(directory: str, filename: str):
    """
    Transcribes an MP3 audio file to text using OpenAI's Whisper model.

    Args:
        mp3_file_path (str): The path to the MP3 audio file.

    Returns:
        str: The transcribed text from the audio file.
    """
    inputfile = os.path.join(directory,  filename + "." + 'mp3')
    outputfile = os.path.join(directory,  filename + "." + 'txt')
    

    
    try:
        # Load the Whisper model (e.g., 'base', 'small', 'medium', 'large')
        # 'base' is a good starting point for general use.
        model = whisper.load_model("base")

        # Transcribe the audio file
        result = model.transcribe(inputfile)

        transcribed_text = result["text"]


        # # Save the transcription to a text file
        # with open(outputfile, "w", encoding="utf-8") as f:
        #     f.write(transcribed_text)

        # print(f"Transcription saved to {outputfile}")
        
        save_transcript_to_docx(transcribed_text,directory,filename)


        # Extract and return the transcribed text
        return result["text"]
    except Exception as e:
        print(f"Error during transcription: {e}")
        return None



def save_transcript_to_docx(transcript_text: str, output_path: str, title: str = None):
    """
    Save a Whisper-transcribed text into a DOCX file.

    :param transcript_text: The transcription text returned by Whisper.
    :param output_path: Path to save the .docx file.
    :param title: Optional title to add at the top of the document.
    """
    doc = Document()

    if title:
        doc.add_heading(title, level=1)

    # Split transcript into paragraphs (you can adapt splitting logic)
    for para in transcript_text.splitlines():
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.save(output_path+"\\"+title+".docx")



def mkv_to_minutes(mkv_path: str, output_docx: str, meeting_title: str = "Meeting Minutes"):
    """
    Convert an MKV meeting video into professional meeting minutes in DOCX.
    
    Steps:
    1. Extract audio from MKV.
    2. Transcribe with Whisper.
    3. Save structured minutes to DOCX.
    """

    # --- 1. Extract audio ---
    audio_path = "temp_audio.mp3"
    print("Extracting audio...")
    video = VideoFileClip(mkv_path)
    video.audio.write_audiofile(audio_path, codec="mp3")

    # --- 2. Transcribe with Whisper ---
    print("Transcribing with Whisper...")
    openai.api_key = os.getenv("OPENAI_API_KEY")
    with open(audio_path, "rb") as f:
        transcript = openai.Audio.transcribe("whisper-1", f)

    transcript_text = transcript["text"]

    # --- 3. Save as Professional Minutes in DOCX ---
    print("Formatting minutes...")
    doc = Document()
    doc.add_heading(meeting_title, level=0)

    doc.add_heading("Date:", level=1)
    doc.add_paragraph("______________________")

    doc.add_heading("Participants:", level=1)
    doc.add_paragraph("______________________")

    doc.add_heading("Agenda:", level=1)
    doc.add_paragraph("______________________")

    doc.add_heading("Discussion:", level=1)
    for para in transcript_text.split(". "):
        if para.strip():
            doc.add_paragraph(f"- {para.strip()}")

    doc.add_heading("Decisions / Action Items:", level=1)
    doc.add_paragraph("______________________")

    doc.save(output_docx)

    # Clean up temp audio
    os.remove(audio_path)
    print(f"Minutes saved to {output_docx}")




placedfiledirectory = "C:\\Users\\Samuel\\Phyton Projects\\Arquivos"
file = "2025-09-19-17-06-01"

# mkv_to_mp3(placedfiledirectory,file)

# transcribe_mp3_to_text(placedfiledirectory,file)

mkv_to_minutes(os.path.join(placedfiledirectory,  file + "." + 'mkv'), file + ".docx",meeting_title="Team Weekly Meeting")